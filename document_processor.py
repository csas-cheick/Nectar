import os
import re
from typing import Dict


class DocumentProcessor:
    """Classe pour extraire et analyser le texte des documents"""
    
    def extract_text(self, filepath: str) -> str:
        """
        Extrait le texte d'un document (PDF, DOCX, TXT)
        """
        extension = os.path.splitext(filepath)[1].lower()
        
        if extension == '.pdf':
            return self._extract_from_pdf(filepath)
        elif extension == '.docx':
            return self._extract_from_docx(filepath)
        elif extension == '.txt':
            return self._extract_from_txt(filepath)
        else:
            raise ValueError(f"Format non supporté: {extension}")
    
    def _extract_from_pdf(self, filepath: str) -> str:
        """Extrait le texte d'un fichier PDF"""
        try:
            import PyPDF2
            
            text = ""
            with open(filepath, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
            
            return text.strip()
        except ImportError:
            # Fallback avec pdfplumber si PyPDF2 échoue
            try:
                import pdfplumber
                
                text = ""
                with pdfplumber.open(filepath) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
                
                return text.strip()
            except ImportError:
                raise ImportError("Installez PyPDF2 ou pdfplumber: pip install PyPDF2 pdfplumber")
    
    def _extract_from_docx(self, filepath: str) -> str:
        """Extrait le texte d'un fichier DOCX"""
        try:
            from docx import Document
            
            doc = Document(filepath)
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extraire aussi le texte des tableaux
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            return text.strip()
        except ImportError:
            raise ImportError("Installez python-docx: pip install python-docx")
    
    def _extract_from_txt(self, filepath: str) -> str:
        """Extrait le texte d'un fichier TXT"""
        encodings = ['utf-8', 'latin-1', 'cp1252']
        
        for encoding in encodings:
            try:
                with open(filepath, 'r', encoding=encoding) as file:
                    return file.read().strip()
            except UnicodeDecodeError:
                continue
        
        raise ValueError("Impossible de décoder le fichier texte")
    
    def get_text_stats(self, text: str) -> Dict:
        """
        Calcule les statistiques du texte
        """
        # Nettoyer le texte pour le comptage
        clean_text = re.sub(r'\s+', ' ', text).strip()
        
        # Compter les mots
        words = clean_text.split()
        word_count = len(words)
        
        # Compter les caractères
        char_count = len(text)
        char_count_no_spaces = len(text.replace(' ', '').replace('\n', ''))
        
        # Compter les phrases (approximatif)
        sentences = re.split(r'[.!?]+', clean_text)
        sentence_count = len([s for s in sentences if s.strip()])
        
        # Compter les paragraphes
        paragraphs = [p for p in text.split('\n\n') if p.strip()]
        paragraph_count = len(paragraphs) if paragraphs else 1
        
        # Temps de lecture estimé (200 mots/minute)
        reading_time_minutes = round(word_count / 200, 1)
        
        return {
            'word_count': word_count,
            'char_count': char_count,
            'char_count_no_spaces': char_count_no_spaces,
            'sentence_count': sentence_count,
            'paragraph_count': paragraph_count,
            'reading_time_minutes': reading_time_minutes
        }
